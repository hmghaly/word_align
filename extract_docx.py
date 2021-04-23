import re, os, json
from itertools import groupby
import base64
#from string import letters, punctuation

from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.text.paragraph import Paragraph
from docx.table import _Cell, Table
from docx.enum.table import WD_TABLE_ALIGNMENT

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
#from docx.enum.style import WD_STYLE_TYPE

def get_docx_paras(docx_fpath):
	document_obj = Document(docx_fpath)
	parent_elm = document_obj.element.body
	all_parapgraphs=[]
#	for child in parent_elm.iterchildren():
#		print child.text
	for child in parent_elm.iterchildren():
		if isinstance(child, CT_P):
			para_obj= Paragraph(child, document_obj)
			all_parapgraphs.append(para_obj)
		elif isinstance(child, CT_Tbl):
			table_obj= Table(child, document_obj)
			for row in table_obj.rows:
				for row in table_obj.rows:
					for cell in row.cells:
						for paragraph in cell.paragraphs:
							all_parapgraphs.append(paragraph)

	return all_parapgraphs

def get_paras_recursive(doc_obj,input_node,all_paras=[],depth=0):
	for child in input_node.iterchildren():
		if isinstance(child, CT_P): 
			para_obj= Paragraph(child, input_node)
			all_paras.append(para_obj)
		else:
			all_paras=get_paras_recursive(doc_obj,child,all_paras,depth+1)

	return all_paras

def get_tables_recursive(doc_obj,input_node,all_tables=[],depth=0):
	for child in input_node.iterchildren():
		#if isinstance(child, CT_P): continue
		#if not isinstance(child, CT_Tbl): continue
		#print child, depth
		if isinstance(child, CT_Tbl): 

			table_obj= Table(child, input_node)
			all_tables.append(table_obj)
			all_tables=get_tables_recursive(doc_obj,child,all_tables,depth+1)
		else:
			all_tables=get_tables_recursive(doc_obj,child,all_tables,depth+1)

	return all_tables




def get_segs(docx_fpath):
	paras=get_docx_paras(docx_fpath)
	all_segs=[]
	for p_i,pa in enumerate(paras):
		el=pa._element
		cur_xml=pa._element.xml
		#xpath_str=str(pa._element.xpath)
		#xpath_id=re.findall("\dx(.+?)>",xpath_str)
		cur_wt_items=get_wt(cur_xml)
		wt_str="".join(cur_wt_items)
		tab_grp=ssplit_full(wt_str)
		#print dir(el)
		#print p_i
		#print el.tag
		#print xpath_id
		for ti, tg in enumerate(tab_grp):
			tg_utf=tg.encode("utf-8","IGNORE")
			par_seg_id=(p_i,ti)
			all_segs.append((par_seg_id,tg_utf))
			#encoded = base64.b64encode(tg.encode("utf-8","IGNORE"))
			#data = base64.b64decode(encoded)
			
			#print p_i,ti, [tg], encoded, [data]

			#print '-----'
		#print tab_grp
		#print '------'
	return all_segs



def write_segs_base64(input_docx_fpath,out_txt_fpath):
	cur_segs=get_segs(input_docx_fpath)
	out_fopen=open(out_txt_fpath,"w")
	for para_seg_id,txt in cur_segs:
		para_seg_id_json=json.dumps(para_seg_id)
		txt_base64=base64.b64encode(txt) 
		line="%s\t%s\n"%(para_seg_id_json,txt_base64)
		out_fopen.write(line)

		#print para_seg_id, para_seg_id_json, [txt], txt_base64
	out_fopen.close()



def ssplit(txt):
    dot_words=["Mr","Ms","Dr","Art","art","Chap","chap","No","no","rev","Rev","Add"] #will need to accomodate preceding space or word boundary
    for dw in dot_words:
        txt=txt.replace(dw+".",dw+"._")
    txt=re.sub("(?u)([\.\?\!\;])\s",r"\1\n",txt) #we need to make sure Arabic, Spanish and Chinese punctuation is also included
    txt=txt.replace("\r","\n")
    txt=txt.replace("\t","\n")
    
    txt=txt.replace("._",".")

    cur_sents=[v.strip() for v in txt.split("\n")]
    cur_sents=[v for v in cur_sents if v]
    return cur_sents

def get_tab_grp(txt): #just to keep the tab and break information
    grouped=[[key,"".join(list(group))] for key,group in groupby(txt,lambda x:x in "\t\n\r")]
    return grouped

def ssplit_full(txt): #keep the tab and break info, while splitting the non-tabs into sentences
    final_list=[]
    tab_grp=get_tab_grp(txt)
    for tg0,tg1 in tab_grp:
        if tg0: final_list.append(tg1)
        else:
            cur_split=ssplit(tg1)
            for cs in cur_split:
                final_list.append(cs)
                final_list.append(" ") #put space items between sentences
            if final_list and final_list[-1]==" ": final_list=final_list[:-1]
    return final_list


def get_wt(txt):
    wt_items=re.finditer('<w:t.*?>(.*?)</w:t>',txt)
    br_items=re.finditer("<w:br/>",txt)
    tab_items=re.finditer("<w:tab/>",txt)
    all_items=[]
    for wt0 in wt_items:
        cur_item=wt0.group(1)
        all_items.append((wt0.start(), cur_item))

    for br0 in br_items:
        all_items.append((br0.start(),"\n"))
    for tab0 in tab_items:
        all_items.append((tab0.start(),"\t"))

    all_items.sort(key=lambda x:x[0])
    final_items=[v[1] for v in all_items]

    return final_items


if __name__=="__main__":
	fname="test-docx-copy.docx"
	#modified_docx_path="test-docx-copy-modified.docx"
	modified_docx_path="test-docx-copy.docx"
	#modified_docx_path="docx-output-modified.docx"
	#modified_docx_path="test-new.docx"
	new_modified_docx_path="new_docx-output-modified1.docx"

	#cur_paras=get_docx_paras(fname)
	#cur_segs=get_segs(fname)
	#write_segs_base64(fname,"out_base64.txt")

	#for cp in cur_paras[:50]:
	#	print cp.text
	#	print '------'

	document_obj = Document(modified_docx_path)
	styles = document_obj.styles
	paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
	#for style in paragraph_styles:
	#	print(style.name)

	mystyle = document_obj.styles.add_style('mystyle', WD_STYLE_TYPE.CHARACTER)
	#style = document_obj.styles.add_style('rtl', WD_STYLE_TYPE.PARAGRAPH)
	style = document_obj.styles.add_style('rtl', WD_STYLE_TYPE.CHARACTER) 

	parent_elm = document_obj.element.body

	cur_tables=get_tables_recursive(document_obj,parent_elm,all_tables=[],depth=0)
	# for ct in cur_tables:
	# 	print ct
	# 	#print ct._element.xml
	# 	print ct.alignment
	# 	ct.alignment = WD_TABLE_ALIGNMENT.RIGHT
	# 	print ct.alignment
	# 	#print ct.xml
	# 	print "========"


	cur_paras=get_paras_recursive(document_obj,parent_elm,all_paras=[])


	for p in cur_paras:
		continue
		#try: print p.style.name
		#except: pass
		#print dir(p)
		#try: print p.style
		#except: pass
		cur_xml=p._element.xml
		cur_txt="".join(get_wt(cur_xml))
		

		
		#if "align" in cur_xml:
		# print cur_xml
		# print cur_txt, [cur_txt]
		# print "======="
		
		only_letters="".join(re.findall("(?u)\w+",p.text))
		english_letters=[v for v in only_letters if v in letters]
		en_ratio=0
		if only_letters: en_ratio=float(len(english_letters))/len(only_letters)
		if en_ratio>0.5:
			p.alignment = WD_ALIGN_PARAGRAPH.LEFT
			p.text=cur_txt
		else:
			p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
			p.clear()
			en_grouped=[(key,list(group)) for key,group in groupby(cur_txt,lambda x:ord(x)<1000 and x!=" ")]
			for e0 in en_grouped:
				status,cluster=e0[0],"".join(e0[1])
				letter_check=False
				if re.findall("[A-Za-z]+",cluster): letter_check=True
				#print status, letter_check, cluster
				#run=p.add_run()
				#run.text=cluster
				# try:
				# 	#, style = "rtl"
				# 	#run=p.add_run(cluster, style = "rtl")
				# 	run=p.add_run(cluster)
				# 	#run.bold = True
				# 	run.rtl=True

				# 	#pass
				# 	#run.style = style
				# 	#font = run.font
				# 	#font.rtl = True
				# except Exception,e: 
				# 	pass
				# 	# print str(e)
			


				
			# print "----"
			# #for ci,c in enumerate(cur_txt):
			# #	print ci, c, ord(c)
			# #en_parts=re.findall("[a-zA-Z\-\/]+",cur_txt)
			# print ">>>", cur_txt
			# #print "en_parts", en_parts


			



		#print p.alignment, [p.text], only_letters, [only_letters], en_ratio

		for run in p.runs:
			continue
			#pass
			run.style = mystyle
			font = run.font
			font.rtl = True

			#print [run.text]

	document_obj.save(new_modified_docx_path)
		#print [p.text]
	#header = document_obj.sections[0].header
	#print header.text
	#print dir(document_obj.sections[0]._sectPr)
	#print document_obj.sections[0]._sectPr.xml
	#print dir(document_obj.sections[0])
	#for a in document_obj.sections[0].iterchildren():
	#	print a

	#for a in document_obj.sections[0]._sectPr.iterchildren():
	#	print a
	#headers = [x.blob.decode() for x in document_obj.part.package.parts if x.partname.find('header')>0]
	#for h in headers:
	#	continue
	#	print h
	#	print "----"


	#for a in cur_segs:
	#	print a
	#for cp0,cp1 in cur_paras:
	#	if cp0=="para": continue
	#	print cp0
	#	#print cp1.text
	#	xpath_str=str(cp1._element.xpath)
	#	xpath_id=re.findall("\dx(.+?)>",xpath_str)
	#	print xpath_str
	#	print xpath_id
	#	print cp1.text
	#	print '-------'